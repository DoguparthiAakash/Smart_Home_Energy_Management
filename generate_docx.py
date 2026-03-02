import os
import re
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: 'python-docx' library not found. Please run 'pip install python-docx' first.")
    exit(1)

def convert_md_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    doc = Document()

    # Title Page
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    title = ""
    if lines and lines[0].startswith("# "):
        title = lines[0].replace("# ", "").strip()
        lines = lines[1:]

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(24)
        doc.add_page_break()

    # Content
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line == "---" or line == "***":
            doc.add_page_break()
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        
        # Lists
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", ""), style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
        
        # Normal text
        else:
            # Handle bold text in line
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part.replace("**", "")).bold = True
                else:
                    p.add_run(part)

    doc.save(docx_file)
    print(f"Successfully generated {docx_file}")

if __name__ == "__main__":
    convert_md_to_docx("Project_Review_Documentation.md", "Project_Review_Documentation.docx")
